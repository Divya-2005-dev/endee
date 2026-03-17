import PyPDF2
import os
import docx
import markdown
from bs4 import BeautifulSoup
from typing import List, Dict, Any
import re
from datetime import datetime

def load_pdf(file_path: str) -> str:
    """
    Load a PDF file and extract text from all pages.
    """
    if file_path.lower().endswith('.pdf'):
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += page_text + "\n"
                return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to read PDF: {e}")
    elif file_path.lower().endswith('.txt'):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError("Unable to decode text file")
    elif file_path.lower().endswith('.docx'):
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to read DOCX: {e}")
    elif file_path.lower().endswith('.md'):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
                # Convert markdown to plain text
                html = markdown.markdown(md_content)
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text()
        except Exception as e:
            raise ValueError(f"Failed to read Markdown: {e}")
    elif file_path.lower().endswith('.html'):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                return soup.get_text()
        except Exception as e:
            raise ValueError(f"Failed to read HTML: {e}")
    else:
        raise ValueError(f"Unsupported file type: {file_path}")

def extract_metadata(file_path: str, filename: str) -> Dict[str, Any]:
    """
    Extract metadata from the file.
    """
    metadata = {
        "filename": filename,
        "file_size": os.path.getsize(file_path),
        "file_type": filename.split('.')[-1].upper() if '.' in filename else "UNKNOWN",
        "uploaded_at": datetime.now().isoformat(),
        "word_count": 0,
        "page_count": 0
    }

    try:
        if file_path.lower().endswith('.pdf'):
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["page_count"] = len(pdf_reader.pages)
                # Try to get PDF metadata
                if pdf_reader.metadata:
                    if pdf_reader.metadata.title:
                        metadata["title"] = pdf_reader.metadata.title
                    if pdf_reader.metadata.author:
                        metadata["author"] = pdf_reader.metadata.author
        elif file_path.lower().endswith('.docx'):
            doc = docx.Document(file_path)
            metadata["page_count"] = len(doc.sections)
            # Try to get document properties
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()

        # Count words in text
        text = load_pdf(file_path)
        metadata["word_count"] = len(text.split())

    except Exception as e:
        metadata["extraction_error"] = str(e)

    return metadata

def split_text_into_chunks(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """
    Split the text into chunks of specified size with overlap.
    """
    if not text.strip():
        return []

    words = text.split()
    if len(words) <= chunk_size:
        return [' '.join(words)]

    chunks = []
    start = 0

    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk = ' '.join(words[start:end])
        chunks.append(chunk)

        # Move start position with overlap
        start = end - overlap
        if start >= len(words):
            break

    return chunks

def clean_text(text: str) -> str:
    """
    Clean and normalize text.
    """
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove non-printable characters
    text = re.sub(r'[^\x20-\x7E\n]', '', text)
    return text.strip()

def detect_language(text: str) -> str:
    """
    Simple language detection (can be enhanced with langdetect library).
    """
    # Basic heuristics
    if any(word in text.lower() for word in ['the', 'and', 'is', 'in', 'to']):
        return 'en'
    elif any(word in text.lower() for word in ['el', 'la', 'de', 'que', 'y']):
        return 'es'
    elif any(word in text.lower() for word in ['le', 'la', 'de', 'et', 'à']):
        return 'fr'
    else:
        return 'unknown'